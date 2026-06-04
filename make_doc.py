from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph style defaults
def style_normal(para, size=11, color=None, bold=False, space_before=0, space_after=6):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        run.font.name = '黑体'
        run.font.color.rgb = RGBColor(0x2b, 0x33, 0x44) if level==1 else RGBColor(0x4f, 0x6f, 0x9e)
    return p

def add_para(doc, text, size=11, bold=False, indent=False, color=None, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '宋体'
    return p

def add_table_row(table, col1, col2, header=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, [col1, col2])):
        cell.text = text
        p = cell.paragraphs[0]
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.font.bold = header
        # shading
        if header:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '4F6F9E')
            tcPr.append(shd)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

# ════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run('PRISM FDI审查机制数据平台')
r.font.size = Pt(22)
r.font.bold = True
r.font.name = '黑体'
r.font.color.rgb = RGBColor(0xBF, 0x6A, 0x4A)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run('功能说明文档')
r2.font.size = Pt(16)
r2.font.bold = True
r2.font.name = '黑体'
r2.font.color.rgb = RGBColor(0x2b, 0x33, 0x44)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(4)
r3 = p3.add_run('基于38国2007–2023年外商直接投资审查机制数据')
r3.font.size = Pt(12)
r3.font.name = '宋体'
r3.font.color.rgb = RGBColor(0x5a, 0x64, 0x77)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('https://jonebroom.github.io/prism-fdi/')
r4.font.size = Pt(11)
r4.font.name = 'Courier New'
r4.font.color.rgb = RGBColor(0x4F, 0x6F, 0x9E)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════
# 1. 平台概述
# ════════════════════════════════════════════
add_heading(doc, '一、平台概述', 1)

add_para(doc, 'PRISM（Policy Review of Investment Screening Mechanisms）是一个专注于外商直接投资（FDI）审查机制的交互式数据可视化平台，系统整理并可视化了全球主要经济体的投资审查制度设计、程序特征与历史演变。', space_after=6)

add_para(doc, '访问地址：https://jonebroom.github.io/prism-fdi/', bold=True, space_after=8)

add_heading(doc, '数据规模', 2)
bullets = [
    '覆盖国家：38个国家（2007–2023年）',
    '时间序列：650+ 国家-年记录',
    '法规事件：143项法规级事件',
    '立法变动：230条立法变动记录',
    '政策文本：141份原始政策文本',
]
for b in bullets:
    add_bullet(doc, b)

add_heading(doc, '国家分组', 2)
for b in ['OECD成员国（经济合作与发展组织）', 'EU/EEA成员国（欧盟及欧洲经济区）', 'Five Eyes（五眼联盟：美国、英国、澳大利亚、加拿大、新西兰）']:
    add_bullet(doc, b)

add_heading(doc, '平台特点', 2)
for b in [
    '全静态网页，无需服务器，无需登录，任何设备通过网址直接访问',
    '所有图表联动：切换年份或国家分组后，全部可视化同步更新',
    '集成AI数据问答功能，支持主流大模型API',
    '内嵌全文检索，覆盖国家名称、法规标题、主管机构、政策文本关键词',
]:
    add_bullet(doc, b)

# ════════════════════════════════════════════
# 2. 左侧控制栏
# ════════════════════════════════════════════
add_heading(doc, '二、左侧控制栏', 1)

add_heading(doc, '年份控制', 2)
add_para(doc, '年份滑块可在2007–2023年间自由拖动，选定年份后所有图表立即联动刷新。点击"▶ Play Timeline"按钮可自动逐年播放，动态观察各国政策的演变过程。', space_after=6)

add_heading(doc, 'Country Groups（国家分组筛选）', 2)
add_para(doc, '三个分组按钮（OECD / EU/EEA / Five Eyes）可单独切换，支持多选组合。取消某分组后，该分组国家从所有图表和列表中移除；至少保留一个分组。', space_after=6)

add_heading(doc, '国家列表', 2)
add_para(doc, '显示当前分组筛选下的全部国家，点击任意国家名称可直接跳转至该国详情页（Country标签）。顶部搜索框支持按国家名称快速过滤。', space_after=6)

# ════════════════════════════════════════════
# 3. 六大模块
# ════════════════════════════════════════════
add_heading(doc, '三、六大分析模块详解', 1)

# ── Tab 01
add_heading(doc, 'Tab 01 · Overview（全球概览）', 2)
add_para(doc, '这是进入平台后的默认页面，提供全球宏观视角下的FDI审查机制全景。', space_after=5)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run('① 世界地图')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
add_para(doc, '按当前年份以颜色深浅显示各国审查机制的某一维度（如严格程度、覆盖类型等）。鼠标悬停可查看该国指标数值，点击国家跳转至详情页。', indent=True, space_after=5)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run('② Global Trend（全球趋势图）')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
add_para(doc, '折线图展示拥有某一审查特征的国家数量随年份的变化趋势；柱状图显示全球每年新增机制数量。图中两条彩色虚线标注了两个关键节点：2019年欧盟FDI法规（蓝色）和2020年新冠疫情引发的立法浪潮（橙色）。', indent=True, space_after=5)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
r = p.add_run('③ Snapshot（快照）')
r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
add_para(doc, '当前年份下各国关键指标的汇总列表，快速纵览所有国家的核心数据。', indent=True, space_after=8)

# ── Tab 02
add_heading(doc, 'Tab 02 · Country（国家详情）', 2)
add_para(doc, '选中某个国家后进入该页，展示该国在当前年份下的全方位数据画像。', space_after=5)

items = [
    ('头部信息', '国家名称、所属分组、机制数量、严格程度评分（0–13分）、覆盖行业数，一目了然。'),
    ('KPI指标卡', '主管机构、覆盖类型、通知要求、预审批要求、审查门槛（股权比例）、审查期限（天数）等核心字段。'),
    ('程序特征雷达图', '将该国的8项核心程序特征与OECD平均水平对比展示；下拉菜单可叠加历史年份进行纵向比较，观察政策收紧或放松的轨迹。'),
    ('功能标志矩阵', '13项程序性权力的逐项打勾显示，包括：正式召回权、跨部门审查、国家安全测试、净收益测试、竞争测试、分级授权、罚款、缓解措施、申报费、本地代表要求、共置要求、增强政府控制、绿地投资覆盖。'),
    ('法规时间线', '该国所有相关法规按时间排列，点击任意一条可在右侧抽屉中查看法规完整详情和原始政策文本。'),
]
for title, desc in items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'· {title}：')
    r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = '宋体'

doc.add_paragraph()

# ── Tab 03
add_heading(doc, 'Tab 03 · Comparison（多国对比）', 2)
add_para(doc, '在同一年份下，将所有国家的多维指标并排展示，适合跨国横向比较。', space_after=5)
for b in [
    '平行坐标图：每条折线代表一个国家，坐标轴包括严格程度、审查门槛、审查期限、行业覆盖数等维度',
    '可拖动各坐标轴上的区间滑块，筛选出满足特定条件的国家子集',
    '鼠标悬停某条线时高亮该国全貌；点击可跳转至该国详情页',
]:
    add_bullet(doc, b)

doc.add_paragraph()

# ── Tab 04
add_heading(doc, 'Tab 04 · Evolution（法规演变）', 2)
add_para(doc, '从多个维度呈现全球FDI审查机制的历史演变规律，包含六个子图表。', space_after=5)

sub_items = [
    ('覆盖类型分布饼图', '当前年份下各覆盖类型（行业型、跨行业型、混合型、资产型等）的国家数量分布，点击扇形可查看该类型所有国家列表'),
    ('Notification × Pre-approval矩阵', '气泡图展示"通知要求"与"预审批要求"两个维度的组合分布，气泡大小代表该组合下的国家数量，点击气泡可查看详情'),
    ('覆盖类型演变流图', '堆叠面积图展示各覆盖类型在2007–2023年间的国家数量变化，直观观察政策模式的全球扩散趋势'),
    ('严格程度分布图', '历年严格程度评分的均值与区间分布，显示全球整体收紧或松动的趋势'),
    ('阈值/时限散点图', '各国机制在审查门槛（股权比例触发线）和审查时限两个维度上的散点分布，可切换横轴指标'),
    ('法规替代关系网络图', '有向网络图展示哪些旧法规被新法规所替代，节点颜色代表国家分组，点击节点可查看国家详情'),
]
for title, desc in sub_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'· {title}：')
    r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = '宋体'

doc.add_paragraph()

# ── Tab 05
add_heading(doc, 'Tab 05 · Changes（立法活动）', 2)
add_para(doc, '聚焦各年度的立法事件，适合追踪特定年份（如2019、2020年）的政策动向。', space_after=5)
for b in [
    '全球立法活动堆叠柱状图：按年份显示新立法、修正案、行政命令、实施细则四类立法活动的数量，点击某年的柱体即可在下方显示该年度所有事件列表',
    '立法事件列表：显示每条事件的国家、法规名称、类型标签（New Law / Amendment / Exec. Order）和摘要说明，点击任意事件可弹出完整政策文本',
]:
    add_bullet(doc, b)

doc.add_paragraph()

# ── Tab 06
add_heading(doc, 'Tab 06 · Sectors（行业分析）', 2)
add_para(doc, '专注于行业层面的覆盖情况分析，适合研究特定行业（如人工智能、半导体、电信）的全球监管扩散。', space_after=5)
for b in [
    '行业覆盖热力图：行（国家）× 列（行业）的矩阵，有颜色代表该国覆盖该行业，支持按行业大类（防务、关键基础设施、新技术等）和细分行业切换',
    '行业排名动态赛马图：可播放逐年动画，展示各行业被覆盖的国家数量随年份竞争增长的动态，直观呈现哪些行业监管扩散最快',
]:
    add_bullet(doc, b)

# ════════════════════════════════════════════
# 4. AI功能
# ════════════════════════════════════════════
add_heading(doc, '四、AI数据问答功能', 1)

add_para(doc, '点击右上角"Ask AI"按钮打开AI对话面板，可用自然语言直接提问，系统基于当前筛选的数据上下文生成回答。', space_after=6)

add_heading(doc, '支持的AI模型', 2)
for b in ['Anthropic Claude（claude-sonnet、claude-haiku等）', 'OpenAI GPT（gpt-4o、gpt-4o-mini等）', 'DeepSeek（deepseek-chat）', 'Moonshot Kimi（moonshot-v1系列）', '智谱GLM（glm-4-flash等）', '任意OpenAI接口兼容的自定义模型']:
    add_bullet(doc, b)

add_heading(doc, '配置方法', 2)
steps = [
    '点击对话框右上角的 ⚙ 图标',
    '在下拉菜单中选择API提供商',
    '填入模型名称和API Key',
    '按 Enter 键或点击"Save"按钮保存',
    '配置保存在浏览器本地，刷新页面不会丢失',
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(f'{i}. {s}')
    run.font.size = Pt(11); run.font.name = '宋体'

add_heading(doc, '问答能力', 2)
for b in [
    '比较两国或多国的审查机制差异（如"比较美国和德国的FDI审查制度"）',
    '识别趋势（如"2020年后哪些国家新增了预审批要求"）',
    '解释字段含义（如"严格程度评分是如何计算的"）',
    '分析特定行业（如"哪些国家将人工智能纳入审查范围"）',
    '页面会根据当前所在标签页和选中国家自动生成4个推荐问题',
]:
    add_bullet(doc, b)

# ════════════════════════════════════════════
# 5. 其他功能
# ════════════════════════════════════════════
add_heading(doc, '五、其他功能', 1)

features = [
    ('全局搜索', '顶部搜索框（快捷键"/"激活），可检索国家名称、法规标题、主管机构名称以及141份政策文本的完整内容。'),
    ('法规详情抽屉', '在Country或Changes页面点击任意法规事件，页面右侧滑出抽屉面板，显示法规完整详情，包括覆盖类型、签署年份、主管机构、政策摘要，以及原始政策文本链接（PDF/HTML）。'),
    ('响应式布局', '支持宽屏显示器、笔记本电脑、平板等不同屏幕尺寸，图表自动适配。'),
    ('键盘快捷键', '"/"键快速激活搜索框，"Escape"键关闭当前面板。'),
]
for title, desc in features:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'· {title}：')
    r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = '宋体'

# ════════════════════════════════════════════
# 6. 数据说明
# ════════════════════════════════════════════
add_heading(doc, '六、数据说明', 1)

add_heading(doc, '数据来源', 2)
add_para(doc, '本平台使用的数据集为 PRISM ISM Dataset（2023.12版本），包含截至2023年底各国FDI审查机制的完整状态及历史演变记录。', space_after=6)

add_heading(doc, '严格程度评分（Strictness Score，0–13分）', 2)
add_para(doc, '由以下13项程序性权力逐项累加得出（具备该项权力计1分，否则计0分）：', space_after=4)
proc = [
    'Formal Call-in（正式召回权）',
    'Review of Increased Ownership（股权增持审查）',
    'Filing Fees（申报费）',
    'Mitigation（缓解措施权）',
    'Fines（罚款权）',
    'National Security Test（国家安全测试）',
    'Net Benefit Test（净收益测试）',
    'Competition Test（竞争测试）',
    'Interagency Review（跨部门联合审查）',
    'Tiered Authority（分级授权）',
    'Local Representation（本地代表要求）',
    'Enhanced Government Control（增强政府控制权）',
    'Co-location（共置要求）',
]
for item in proc:
    add_bullet(doc, item, size=10)

add_heading(doc, '覆盖类型分类', 2)
coverage = [
    ('Sectoral（行业型）', '仅针对特定敏感行业实施审查'),
    ('Cross-sectoral（跨行业型）', '对所有或大多数行业实施统一审查'),
    ('Mixed（混合型）', '兼具行业型和跨行业型特征'),
    ('Asset-based（资产型）', '以资产类别（如土地、基础设施）为触发条件'),
]
for ctype, desc in coverage:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(f'{ctype}：')
    r.font.bold = True; r.font.size = Pt(11); r.font.name = '宋体'
    r2 = p.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = '宋体'

# ════════════════════════════════════════════
# Save
# ════════════════════════════════════════════
out = r'D:\社科数据分析的项目\投资审查网页制作\PRISM平台功能说明.docx'
doc.save(out)
print('Saved:', out)
